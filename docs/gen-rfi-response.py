"""Generate OPM OIG RFI Response Word Document."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_body(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")

# ═══════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Y Point Technologies"); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OIG-ITS"); r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Investigative Tracking System"); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x33, 0x44, 0x55)

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Response to Request for Information"); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x25, 0x63, 0xeb); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Office of Personnel Management\nOffice of Inspector General"); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x33, 0x44, 0x55)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("April 2026"); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FOUO \u2014 For Official Use Only"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26); r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. Y Point\u2019s Recommended Solution: Custom-Built ITS on AWS GovCloud",
    "   1.1 Solution Architecture Overview",
    "   1.2 Current Operational Capabilities",
    "   1.3 NL2SQL Application",
    "   1.4 AI/ML Pipeline",
    "2. Custom ITS Technical Architecture",
    "   2.1 Database Layer",
    "   2.2 Application Layer",
    "   2.3 AI/ML Integration Layer",
    "   2.4 GenAI Case Management Capabilities",
    "   2.5 Security Architecture",
    "   2.6 Authentication: Okta PIV-Based Login",
    "3. SOW Requirement Compliance",
    "   3.1 Security, Compliance, and Authorization",
    "   3.2 Workflow, Automation, and Business Rules Engine",
    "   3.3 Case Management Capabilities",
    "   3.4 Document and Records Management",
    "   3.5 External Intake and Portal Capabilities",
    "   3.6 Reporting and Analytics",
    "   3.7 Training, Timekeeping, and Supporting Modules",
    "   3.8 Performance, Scalability, and SLAs",
    "   3.9 Administration and Configuration",
    "4. Data Integration and Migration",
    "   4.1 Native Data Warehouse Integration",
    "   4.2 Legacy Data Migration",
    "   4.3 External System Integration",
    "5. AI Differentiators: Capabilities Not Available in Leading Industry Solutions",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECTION 1
# ═══════════════════════════════════════════════════════
doc.add_heading("1. Y Point\u2019s Recommended Solution: Custom-Built ITS on AWS GovCloud", level=1)

doc.add_heading("1.1 Solution Architecture Overview", level=2)
add_body(
    "Y Point recommends building the Investigative Tracking System as a custom application deployed on top of "
    "the existing OPM OIG AWS GovCloud FedRAMP High infrastructure that Y Point has already designed and operates. "
    "This solution leverages the foundational data assets and AI/ML capabilities that are already in production or "
    "under active development:"
)
add_bullet("Master Data: Comprehensive provider, member, and drug master data for unified entity resolution\u2014already operational.")
add_bullet("Reference and External Data: Code descriptions, carrier benefit brochures, conflict of interest data, and DOJ case references\u2014under active curation.")
add_bullet("Standardized Claims: Common claims data model with detailed data dictionary for consistent analytics across all FEHBP carriers\u2014already operational.")
add_bullet("GenAI Explanation and Document Creation: Automated claim explanation, cluster description, patient profiling, and audit rationale generation\u2014under active development as part of the FY2027 AI Strategic Plan, leveraging AWS GovCloud Bedrock foundation models.")

doc.add_heading("1.2 Current Operational Capabilities", level=2)
add_body(
    "Y Point has envisioned a SchemeRadar Investigator Board, an AI/ML application that analyzes billions of "
    "healthcare claims and presents ranked clusters of anomalous activity to OIG investigators. Each cluster card "
    "displays the dollar-at-risk, claim counts, fraud probability score, trend visualization, identified motifs, "
    "and quick-action buttons for viewing evidence or starting a case."
)
add_body(
    "In the planned architecture, investigators will be able to select anomalous claim clusters directly from this "
    "board and create fully populated investigative cases with a single click. The case record will be pre-populated "
    "with subject/provider profiles, claim evidence, anomaly scores, financial calculations, timeline data, and "
    "GenAI-generated rationale\u2014eliminating hours of manual data entry per case."
)

doc.add_heading("1.3 NL2SQL Application", level=2)
add_body(
    "Y Point is in the process of implementing a production Natural Language to SQL (NL2SQL) application that allows "
    "OIG investigators and auditors to query the data warehouse using plain English. Investigators can ask questions "
    'such as "Show me all claims for provider X in Q3 2025 where the paid amount exceeded expected by more than 50%" '
    "and receive structured query results instantly. This capability will be directly integrated into the custom ITS, "
    "enabling investigators to perform ad hoc data exploration without leaving the case management environment."
)

doc.add_heading("1.4 AI/ML Pipeline", level=2)
add_body(
    "The AI/ML pipeline processes billions of claims records and applies anomaly scoring algorithms, pattern recognition "
    "models, and clustering techniques to identify claim groups that warrant investigative attention. The FY2027 AI "
    "Strategic Plan expands this foundation with LSTM time-series analysis for detecting emerging fraud patterns, "
    "custom claim embedding models for similarity matching, and XGBoost supervised classification trained on improved "
    "anomaly scores."
)

# ═══════════════════════════════════════════════════════
# SECTION 2
# ═══════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("2. Custom ITS Technical Architecture", level=1)

doc.add_heading("2.1 Database Layer", level=2)
add_body(
    "The custom ITS will be built on a PostgreSQL database deployed within the existing AWS GovCloud infrastructure. "
    "Y Point has already constructed a fully functional case management database schema and application prototype using "
    "PostgreSQL with 53 data models, 100+ indexes, and comprehensive relationship mapping. This database is designed to "
    "operate natively alongside the DMG data warehouse, sharing entity resolution tables, reference data, and claims data "
    "without requiring ETL translation layers. The system can be deployed into the current AWS GovCloud environment immediately."
)

doc.add_heading("2.2 Application Layer", level=2)
add_body(
    "The web application will be a modern, browser-based interface built on Next.js and React, requiring no local client "
    "installations, plugins, or legacy browser dependencies\u2014fully compliant with the SOW\u2019s browser-only mandate. "
    "The application will be designed with role-based access controls at both the case and field level, supporting "
    "SAML/OIDC integration for single sign-on and multi-factor authentication through the OIG\u2019s existing Okta identity "
    "provider infrastructure with PIV card authentication."
)

doc.add_heading("2.3 AI/ML Integration Layer", level=2)
add_body(
    "Unlike any COTS solution, the custom ITS will feature native, bidirectional integration with Y Point\u2019s AI/ML "
    "pipeline. This means that anomaly-detected claim clusters flow directly into the case management workflow; investigator "
    "actions and case outcomes feed back into the ML models to improve future detection accuracy; and GenAI modules can "
    "automatically generate case summaries, participant profiles, event timelines, overpayment calculations, brochure "
    "citations, and draft report sections directly within the case record."
)
add_body(
    "All GenAI capabilities are powered through AWS GovCloud Bedrock foundation models. The architecture is model-agnostic "
    "and swappable\u2014as newer, more advanced foundation models become available on AWS GovCloud Bedrock, the application "
    "can adopt them without code changes, ensuring the system continuously benefits from advances in AI technology."
)

doc.add_heading("2.4 GenAI Case Management Capabilities", level=2)
add_body(
    "The custom ITS incorporates ten GenAI-powered capabilities specifically designed for investigative case management, "
    "deployed on AWS GovCloud Bedrock. These capabilities are integrated directly into the case management workflow and "
    "operate on case data in real time:"
)
genai = [
    ("Case Similarity Analysis", "Cosine similarity on 22-dimensional feature vectors (case type, priority, complaint source, subject count, evidence count, financial totals). Identifies the most similar historical cases to inform investigation strategy. Investigators can see which past cases had similar profiles and what outcomes resulted."),
    ("Case Clustering", "K-means clustering (k=5) groups active investigations by attribute similarity. Clusters are labeled by dominant case type and complaint source. Enables supervisors to identify patterns across their portfolio and assign resources to related case groups."),
    ("Document Auto-Classification", "Every uploaded document is automatically classified into one of eight categories (legal/subpoena, interview memo, investigation report, financial record, photographic evidence, correspondence, contract, court document) with confidence scoring. Entity extraction identifies dollar amounts, dates, case numbers, and agency names. Classification runs on upload with zero investigator effort."),
    ("Investigator Recommendation Engine", "When a new case is created, the system recommends the optimal investigator based on four weighted factors: expertise match (30 points\u2014ratio of past cases with same type), current workload (25 points\u2014inverse of active case count), historical success rate (25 points\u2014ratio of substantiated cases), and availability (20 points\u2014no cases due this week). Returns the top three candidates with detailed scoring breakdowns."),
    ("Evidence Strength Scoring", "Scores each case\u2019s evidence portfolio 0\u2013100 with a letter grade (A through F). Scoring factors include per-type item points (physical evidence +5, digital +4, documentary +3, testimony +2), chain-of-custody completeness (+10), evidence type diversity (+15 for 3+ types), source corroboration (+10 for multiple independent sources), and timeline coverage (+5 for evidence spanning 30+ days)."),
    ("Workload Balancing Analysis", "Calculates a weighted workload score per investigator: active cases (\u00d73), critical cases (\u00d75), pending tasks (\u00d71), overdue tasks (\u00d74), and cases due this week (\u00d73). Identifies overloaded investigators (score > mean + 1.5\u03c3) and underloaded investigators (score < mean \u2013 1\u03c3). Supervisor approval queue depth is also tracked."),
    ("Subject Risk Profiling", "Scores each subject (individual, organization, or vendor) 0\u2013100 based on: number of cases involved (+5 per case), violations (+10 each, +20 if substantiated), financial impact (logarithmic scale from $1K to $1M+), network hub connections, repeat offender status (+15 if involved in 3+ cases), and role escalation patterns (+10 if upgraded from witness to respondent across investigations)."),
    ("Case Narrative Generation", "Automatically generates structured, plain-English investigation summaries from case data. The narrative includes sections for case opening (case number, type, opened date, subject count), subjects and their roles, investigative techniques employed with dates and findings, documented violations with status and disposition, financial impact with per-subject breakdowns, agency referrals with outcomes, and current case status with closure readiness assessment."),
    ("Interview Question Generation", "Given the case type, subject role (complainant, respondent, witness), case description, and known facts, the system generates targeted interview questions organized into four categories: opening questions to establish rapport and baseline, substantive questions addressing the specific allegations, probe questions for follow-up based on anticipated responses, and closing questions for documentation and next steps. Interview tips specific to the case type are also provided."),
    ("Intelligent Report Generation", "Generates professional investigation reports from complete case data. Supports four report types: summary (executive overview), narrative (detailed chronological account), findings (structured evidentiary analysis with conclusions), and recommendation (action items with supporting rationale). The system fetches all case relations\u2014subjects, violations, financials, techniques, referrals, notes, evidence, assignments\u2014and produces publication-ready documents."),
]
for title, desc in genai:
    p = doc.add_paragraph()
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(desc)

add_body(
    "All ten GenAI capabilities operate through AWS GovCloud Bedrock\u2019s model-agnostic API layer. The underlying "
    "foundation models are swappable without application code changes. As AWS makes newer, more capable models available "
    "on GovCloud Bedrock, the ITS will automatically benefit from improved accuracy, speed, and capability\u2014a future-proofing "
    "advantage that no COTS solution can match."
)

doc.add_heading("2.5 Security Architecture", level=2)
add_body(
    "The solution inherits the FedRAMP High authorization of the existing AWS GovCloud infrastructure. All data at rest "
    "is encrypted using AES-256, and all data in transit is protected by TLS 1.2/1.3. The system supports granular "
    "role-based access controls with six defined roles and 35+ discrete permissions enforced at both the case and field "
    "level, comprehensive audit logging with field-level change tracking (old and new values) exportable to Microsoft "
    "Defender for Cloud Apps via SFTP, and HIPAA-compliant handling of electronic Protected Health Information. Backup "
    "and disaster recovery will leverage AWS GovCloud\u2019s multi-region capabilities, with an alternate processing site "
    "located outside the 100-mile radius of Washington, D.C. as specified in the SOW."
)

doc.add_heading("2.6 Authentication: Okta PIV-Based Login", level=2)
add_body(
    "The ITS will integrate with the OIG\u2019s existing Okta identity provider for authentication, supporting PIV "
    "(Personal Identity Verification) card-based login in compliance with HSPD-12 and FIPS 201. The authentication "
    "architecture supports:"
)
add_bullet("PIV card authentication via Okta\u2019s Smart Card IDP integration, enabling certificate-based login using government-issued PIV credentials.")
add_bullet("SAML 2.0 and OIDC protocol support for seamless single sign-on (SSO) with the OIG\u2019s Okta tenant.")
add_bullet("Multi-factor authentication (MFA) enforced through Okta\u2019s adaptive MFA policies, including PIV + PIN as a phishing-resistant factor.")
add_bullet("Session management with configurable timeout (default 8 hours), account lockout after 5 failed attempts (15-minute cooling period), and one-session-per-user enforcement.")
add_bullet("Role synchronization: user roles and group memberships defined in Okta are automatically mapped to the ITS\u2019s six-role RBAC model (Administrator, Supervisor, Investigator, Analyst, Auditor, Read-Only) upon login.")
add_bullet("Automated provisioning and deprovisioning: when users are added or removed in Okta, their ITS access is updated in real time via SCIM or Okta lifecycle hooks.")
add_body(
    "This approach eliminates the need for local password management within the ITS. All credential lifecycle "
    "management\u2014issuance, rotation, revocation\u2014is handled by Okta, reducing the application\u2019s attack "
    "surface and ensuring compliance with OMB M-22-09 zero trust mandates."
)

# ═══════════════════════════════════════════════════════
# SECTION 3
# ═══════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("3. SOW Requirement Compliance", level=1)

doc.add_heading("3.1 Security, Compliance, and Authorization", level=2)
add_body(
    "The system will operate within the existing OPM OIG AWS GovCloud environment, which maintains FedRAMP High "
    "authorization. The infrastructure is FISMA compliant at the Moderate-to-High impact level, implements NIST SP "
    "800-53 security and privacy controls, and is FIPS 199/200 compliant. Y Point has extensive experience preparing "
    "Authority to Operate (ATO) artifacts including System Security Plans (SSP), Plans of Action and Milestones "
    "(POA&M), and security assessment packages. The architecture fully supports HIPAA Security and Breach Notification "
    "Rules for the protection of ePHI associated with FEHBP investigations. Identity and access management will "
    "integrate with the OIG\u2019s Okta identity provider using SAML 2.0 and/or OIDC protocols with PIV card "
    "authentication, providing SSO and MFA capabilities. All audit logs will capture user identity, timestamp, and "
    "the specific data element modified, with native export to Microsoft Defender for Cloud Apps. Incident response "
    "SLAs will include notification of the OIG system administrator within one hour of discovering a suspected cyber incident."
)

doc.add_heading("3.2 Workflow, Automation, and Business Rules Engine", level=2)
add_body(
    "The custom application will include an event-driven business rules engine configurable through an administrative "
    "interface without requiring code changes. Workflows will support automated task routing and escalation, deadline "
    "and SLA tracking with visual urgency indicators (badges and color-coded countdowns) on user dashboards, automated "
    "email and in-application notifications triggered by configurable events, mandatory field validation preventing case "
    "closure without required data, supervisory approval chains with support for temporary delegation of authority, and "
    "role-based report page linkages to dashboards. Administrators will be able to modify workflow rules, add database "
    "fields, adjust dropdown menus, and configure notification triggers through a front-end administrative "
    "interface\u2014meeting the RFI\u2019s mandate that the system be maintainable by a single FTE."
)

doc.add_heading("3.3 Case Management Capabilities", level=2)
add_body(
    "The system will support the full case lifecycle from intake through disposition, including fluid transitions "
    "between complaint, preliminary inquiry, and full investigation stages with automatic retention of all historical "
    "data. Entity management will support subjects, victims, witnesses, and corporate entities with dynamic status "
    "changes and hierarchical relationship mapping. Cross-case entity resolution will identify shared subjects across "
    "multiple investigations. The system will integrate with the SAM API to auto-populate corporate entity profiles. "
    "Case chronology logs will automatically record all status changes with user attribution and timestamps. Evidence "
    "tracking will maintain a complete chain of custody with automated logging of all evidence-related actions."
)

doc.add_heading("3.4 Document and Records Management", level=2)
add_body(
    "The system will support upload and storage of all standard file formats (30+ types including PDF, Word, Excel, "
    "PowerPoint, images, audio, video, and email) with automated metadata extraction and full-text search capabilities, "
    "including Boolean operators. WebDAV integration will enable online editing of Microsoft Office documents within the "
    "case environment. Electronic signatures will lock approved documents, convert them to searchable PDF, and "
    "automatically purge drafts. If an approved document is subsequently altered, the signature will be automatically "
    "revoked. Records retention schedules will comply with NARA standards, with automated purge/archive rules and legal "
    "hold capabilities. Document versioning tracks all revisions with version numbering (v1, v2, v3). Nineteen "
    "pre-configured document templates auto-populate with case data for common investigative documents including "
    "subpoenas, memoranda of interview, reports of investigation, and evidence receipts."
)

doc.add_heading("3.5 External Intake and Portal Capabilities", level=2)
add_body(
    "The custom solution will include secure, external-facing web portals for hotline complaints, whistleblower "
    "submissions, and FEHBP carrier fraud notifications. All portals will feature encrypted form submissions with "
    "data validation, automated acknowledgment responses with unique tracking identifiers, and status checking "
    "capabilities for submitters. Whistleblower submissions will be automatically segregated with strict "
    "confidentiality protocols and legal protections prominently displayed per the Whistleblower Protection Act "
    "(5 U.S.C. \u00a7 2302(b)(8)). The FEHBP Carrier Notification Portal will validate submissions against OIG "
    "business rules, provide descriptive error messages for rejected entries, support offline caching for interrupted "
    "submissions, and automatically parse received data to create preliminary inquiries or link to existing "
    "investigations using entity resolution."
)

doc.add_heading("3.6 Reporting and Analytics", level=2)
add_body(
    "The system will maintain pre-configured report templates for the Semiannual Report to Congress (SAR) and CIGIE "
    "metrics. The ad hoc reporting interface will allow non-technical staff to construct complex queries using Boolean "
    "operators, date range filters, and point-in-time snapshots without requiring SQL knowledge or vendor intervention. "
    "Results will be exportable to Excel, CSV, and PDF formats. Role-based dashboards will display real-time KPIs with "
    "drill-down capability from aggregate statistics to individual case records. Because the ITS shares the same data "
    "warehouse as the AI/ML analytics pipeline, investigators will have access to analytics that no COTS product could "
    "provide\u2014including anomaly trend visualizations, peer comparison data, and ML-generated risk assessments "
    "directly within their case dashboard."
)

doc.add_heading("3.7 Training, Timekeeping, and Supporting Modules", level=2)
add_body(
    "The system will include dedicated modules for: LEAP time and labor tracking with automated calculation of "
    "unscheduled duty hours, monthly and annual compliance aggregation, case attribution, and supervisory certification "
    "workflows in full compliance with 5 U.S.C. \u00a7 5545a; training management with personalized learning plans, "
    "automated certification expiration notifications, document upload for qualification records, and financial tracking "
    "of training costs; and accountable property and evidence inventory management linking physical assets to staff "
    "members with full chain-of-custody audit trails. Financial calculation capabilities will support custom formulas "
    "for restitution allocation, lost investment income computation, and aggregate financial outcome tracking across "
    "the investigative portfolio."
)

doc.add_heading("3.8 Performance, Scalability, and SLAs", level=2)
add_body(
    "The system will support a minimum of 30 concurrent licensed users with page load times under two seconds for at "
    "least 23 hours per day. The AWS GovCloud infrastructure provides elastic scalability to accommodate growth in "
    "data volume and user count. Uptime targets will meet or exceed 99.5% availability. The existing infrastructure "
    "already processes billions of claims records, demonstrating proven capacity for the data volumes the ITS will manage."
)

doc.add_heading("3.9 Administration and Configuration", level=2)
add_body(
    "Administrative capabilities will include user and role management, field-level configuration (adding fields, "
    "modifying dropdowns, adjusting labels), workflow and notification management, and report template "
    "creation\u2014all through a front-end graphical interface. The level of effort for ongoing configuration will be "
    "minimal, fully supporting the OIG\u2019s one-FTE O&M target. Unlike COTS products that require vendor intervention "
    "for schema changes, the custom system will give OIG administrators direct control over all configurable elements."
)

# ═══════════════════════════════════════════════════════
# SECTION 4
# ═══════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("4. Data Integration and Migration", level=1)

doc.add_heading("4.1 Native Data Warehouse Integration", level=2)
add_body(
    "The most significant architectural advantage of Y Point\u2019s approach is the elimination of the data integration "
    "problem that plagues COTS deployments. Because the custom ITS operates within the same AWS GovCloud environment "
    "and shares the DMG data warehouse\u2019s data structures, there is no need for complex ETL mapping between "
    "incompatible schemas. Case data, entity profiles, claims evidence, and financial calculations all reference the "
    "same underlying data model. This means no data duplication, no synchronization failures, and no ongoing integration maintenance."
)

doc.add_heading("4.2 Legacy Data Migration", level=2)
add_body(
    "Y Point will execute a structured data migration from the incumbent legacy system following a rigorous methodology: "
    "comprehensive data mapping and schema reconciliation between legacy and target databases; a mandatory Proof of "
    "Concept (PoC) migration using a sanitized subset of investigations including entities, chronological logs, and "
    "attached documents; validation that all metadata (creation dates, user attribution, approval timestamps) is "
    "perfectly preserved; and tested rollback strategies for catastrophic data corruption scenarios. The migration will "
    "leverage Y Point\u2019s existing expertise with the OIG\u2019s data environment, significantly reducing the mapping "
    "effort compared to migrating data into an unfamiliar COTS schema."
)

doc.add_heading("4.3 External System Integration", level=2)
add_body(
    "The system will support API-based integration with the SAM API for corporate entity verification, Microsoft Office "
    "products via WebDAV, SIEM tools (Microsoft Defender for Cloud Apps) via SFTP, and the OIG\u2019s Okta identity "
    "provider via SAML/OIDC with PIV card authentication. Y Point\u2019s existing data warehouse already ingests data "
    "from multiple external sources at high volume, demonstrating proven integration capability."
)

# ═══════════════════════════════════════════════════════
# SECTION 5 — AI DIFFERENTIATORS
# ═══════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("5. AI Differentiators: Capabilities Not Available in Leading Industry Solutions", level=1)
add_body(
    "Y Point\u2019s custom ITS includes ten GenAI-powered case management capabilities that are not available in any "
    "of the five leading OIG case management solutions evaluated (eCASE Investigations, CMTS, ICMS with LEAP, "
    "Entellitrak IG App, and Investigative CM by ServiceNow). The following table documents each unique capability "
    "and its absence from competitive offerings."
)

# AI Differentiator table
t = doc.add_table(rows=11, cols=4)
t.style = "Light Grid Accent 1"
headers = ["GenAI Capability", "Y Point ITS", "Available in eCASE / CMTS / ICMS / Entellitrak / ServiceNow?", "Investigative Value"]
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
    for p in t.rows[0].cells[j].paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(8)

diff_rows = [
    ("Case Similarity Analysis", "Yes \u2014 cosine similarity on 22-dim feature vectors", "No \u2014 none of the five solutions offer automated case similarity matching", "Identifies related historical cases to inform investigation strategy and predict outcomes"),
    ("Case Clustering", "Yes \u2014 k-means grouping of active investigations", "No \u2014 not available in any evaluated solution", "Reveals patterns across the investigation portfolio; enables resource allocation to case groups"),
    ("Document Auto-Classification", "Yes \u2014 automatic on every upload with confidence scoring", "No \u2014 eCASE has basic AI search but no auto-classification; others have none", "Eliminates manual document tagging; ensures consistent categorization across investigators"),
    ("Investigator Recommendation", "Yes \u2014 4-factor scoring (expertise, workload, success rate, availability)", "Partial \u2014 eCASE has auto-assignment but not multi-factor recommendation; others have none", "Optimizes case assignment based on data rather than availability alone"),
    ("Evidence Strength Scoring", "Yes \u2014 0-100 score with A-F letter grade across 5 factors", "No \u2014 not available in any evaluated solution", "Quantifies evidence quality; identifies gaps before case closure or prosecution referral"),
    ("Workload Balancing", "Yes \u2014 statistical analysis with overload/underload detection", "No \u2014 not available in any evaluated solution", "Prevents investigator burnout; ensures equitable case distribution"),
    ("Subject Risk Profiling", "Yes \u2014 multi-factor 0-100 risk score per subject", "No \u2014 not available in any evaluated solution", "Prioritizes subjects for investigation based on quantified risk rather than intuition"),
    ("Case Narrative Generation", "Yes \u2014 structured plain-English summaries from case data", "Partial \u2014 ServiceNow has basic AI summarization; others have none", "Reduces hours of report writing to seconds; ensures consistent narrative structure"),
    ("Interview Question Generation", "Yes \u2014 role-specific questions organized by interview phase", "No \u2014 not available in any evaluated solution", "Standardizes interview preparation; ensures comprehensive coverage of case facts"),
    ("Intelligent Report Generation", "Yes \u2014 4 report types (summary, narrative, findings, recommendation)", "No \u2014 ServiceNow has partial AI assist; others have none", "Produces publication-ready investigation reports directly from case data"),
]
for i, (cap, ypoint, avail, value) in enumerate(diff_rows):
    t.rows[i+1].cells[0].text = cap
    t.rows[i+1].cells[1].text = ypoint
    t.rows[i+1].cells[2].text = avail
    t.rows[i+1].cells[3].text = value
    for c in t.rows[i+1].cells:
        for p in c.paragraphs:
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for r in p.runs: r.font.size = Pt(8)

doc.add_paragraph()
add_body(
    "Model Swappability: All ten GenAI capabilities are deployed through AWS GovCloud Bedrock\u2019s model-agnostic "
    "API layer. The underlying foundation models are swappable without application code changes. As AWS makes newer, "
    "more capable foundation models available on GovCloud Bedrock, the ITS will automatically benefit from improved "
    "accuracy, speed, and capability. This future-proofing advantage ensures that the OIG\u2019s investment in AI "
    "capabilities appreciates over time rather than depreciating as competing products\u2019 proprietary models age."
)

add_body(
    "Integration with Existing AI/Analytics: Since the custom ITS is built on top of OPM OIG\u2019s existing "
    "AI and analytics infrastructure operated by Y Point, the ten GenAI capabilities listed above complement\u2014"
    "rather than duplicate\u2014the existing anomaly detection, claims analysis, and fraud scoring capabilities "
    "already in production. The combined platform delivers an investigation intelligence capability that no COTS "
    "solution can replicate."
)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Y Point Technologies  |  April 2026  |  FOUO")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save("docs/OIG-ITS_RFI_Response.docx")
print(f"Created: docs/OIG-ITS_RFI_Response.docx ({os.path.getsize('docs/OIG-ITS_RFI_Response.docx')//1024} KB)")
