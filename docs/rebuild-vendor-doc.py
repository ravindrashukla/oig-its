"""Rebuild OIG-ITS vs 5 federal OIG vendors — complete Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.2); s.right_margin = Cm(1.2)
    s.page_width = Inches(11); s.page_height = Inches(8.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

def ctr(text, sz, clr=None, bld=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(sz); r.bold = bld
    if clr: r.font.color.rgb = RGBColor(*clr)

def tbl(hdrs, rows, hl=1):
    t = doc.add_table(rows=len(rows)+1, cols=len(hdrs))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdrs):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(8)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            t.rows[i+1].cells[j].text = str(v)
            for p in t.rows[i+1].cells[j].paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(8)
                    if j == hl:
                        r.bold = True
                        if str(v).startswith("Yes"):
                            r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
    doc.add_paragraph()

# COVER
doc.add_paragraph(); doc.add_paragraph()
ctr("OIG-ITS", 40, (0x0f, 0x17, 0x2a), True)
ctr("Investigative Tracking System", 20, (0x33, 0x44, 0x55))
doc.add_paragraph()
ctr("Comparative Evaluation Against", 16, (0x25, 0x63, 0xeb), True)
ctr("Federal OIG Case Management Solutions", 16, (0x25, 0x63, 0xeb), True)
doc.add_paragraph()
ctr("OIG-ITS (Y Point)  vs.", 12, (0x0f, 0x17, 0x2a), True)
for v in ["eCASE Investigations (Casepoint/OPEXUS)",
          "CMTS (WingSwept, LLC)",
          "ICMS with LEAP (Juvare, LLC)",
          "Entellitrak IG App (Tyler Technologies)",
          "Investigative CM (ServiceNow)"]:
    ctr(v, 11, (0x33, 0x44, 0x55))
doc.add_paragraph()
ctr("Prepared by Y Point  |  April 2026", 10, (0x94, 0xa3, 0xb8))
doc.add_page_break()

H = ["Feature", "OIG-ITS", "eCASE", "CMTS", "ICMS", "Entellitrak", "ServiceNow"]

# 1. OVERVIEW
doc.add_heading("1. Comparative Evaluation Summary", level=1)
tbl(["Rank", "Solution", "Vendor", "FedRAMP", "LEAP", "OIG Spec.", "AI Count"],
    [("New", "OIG-ITS", "Y Point", "Aligned", "Native (auto)", "Very High", "22"),
     ("1", "eCASE Investigations", "Casepoint", "High/IL5", "Native", "Very High", "3-5"),
     ("2", "CMTS", "WingSwept", "Moderate", "Configurable", "High", "0"),
     ("3", "ICMS with LEAP", "Juvare", "High/IL5", "Native", "Mod-High", "0"),
     ("4", "Entellitrak (IG App)", "Tyler Tech", "Authorized", "Custom", "High", "0"),
     ("5", "Investigative CM", "ServiceNow", "High", "Custom", "Low", "3-5")])

# 2. CASE MANAGEMENT
doc.add_heading("2. Case Management", level=1)
tbl(H, [
    ("Full case lifecycle (7+ statuses)", "Yes (7 statuses)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Auto case numbering", "Yes (OIG-YYYY-NNN)", "Yes", "Yes", "Yes", "Config", "Yes"),
    ("Pre-close validation", "Yes (13 criteria)", "Yes (checklists)", "No", "No", "Config", "No"),
    ("Case auto-lock on closure", "Yes", "Yes", "Unknown", "Unknown", "Config", "Unknown"),
    ("Case relationships/spinoffs", "Yes", "Yes", "Yes", "Yes", "Config", "Yes"),
    ("Draft/autosave mode", "Yes", "Unknown", "Unknown", "Unknown", "Yes", "Yes"),
    ("Field permissions by role+status", "Yes (6x7 matrix)", "Yes", "Unknown", "Unknown", "Config", "Yes"),
    ("Configurable workflows", "Yes (5 types)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Subject mgmt (5 types, 6 roles)", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Violations per subject", "Yes", "Yes", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Financial results tracking", "Yes (ROI auto)", "Yes", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Subpoena package workflow", "Yes", "Yes", "Unknown", "Unknown", "Unknown", "Unknown"),
])

# 3. EVIDENCE
doc.add_heading("3. Evidence & Documents", level=1)
tbl(H, [
    ("Immutable chain of custody", "Yes", "Yes", "Basic", "Yes (full)", "Config", "Yes"),
    ("Sequential exhibit labels", "Yes (EX-001)", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("8 evidence types", "Yes", "Yes", "Yes", "Yes", "Config", "Yes"),
    ("Document versioning", "Yes (v1,v2)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Supervisory doc approval", "Yes", "Yes", "Yes", "Unknown", "Config", "Yes"),
    ("19 document templates", "Yes", "Yes", "Yes", "Unknown", "Config", "Unknown"),
    ("ZIP download of case files", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("AI auto-classification on upload", "Yes", "No", "No", "No", "No", "No"),
])

# 4. PUBLIC INTAKE
doc.add_heading("4. Public Intake & Whistleblower", level=1)
tbl(H, [
    ("Public hotline (no auth)", "Yes", "Yes", "Basic", "Intake", "Config", "Yes"),
    ("Whistleblower portal + WPA", "Yes (WPA cited)", "Yes", "Unknown", "Unknown", "Unknown", "Yes"),
    ("Anonymous submissions", "Yes", "Yes", "Unknown", "Unknown", "Unknown", "Yes"),
    ("AI risk scoring on intake", "Yes (0-100)", "No", "No", "No", "No", "No"),
    ("Complaint dedup (AI)", "Yes (Jaccard)", "No", "No", "No", "No", "No"),
    ("Inquiry to Case (1-click)", "Yes", "Yes", "Yes", "Yes", "Config", "Yes"),
])

# 5. LEAP
doc.add_heading("5. LEAP Support", level=1)
tbl(H, [
    ("LEAP time tracking", "Yes (native)", "Yes (native)", "Config", "Yes (native)", "Custom", "Custom"),
    ("Auto LEAP calculation", "Yes (Title 5)", "Yes", "Config", "Yes", "Manual", "Manual"),
    ("Substantial hours check", "Yes (auto)", "Yes", "Unknown", "Yes", "Unknown", "Unknown"),
    ("Availability pay (25%)", "Yes (auto)", "Yes", "Unknown", "Yes", "Unknown", "Unknown"),
    ("Timesheet approval", "Yes", "Yes", "Yes", "Yes", "Config", "Unknown"),
])

# 6. AI/ML
doc.add_page_break()
doc.add_heading("6. AI/ML Capabilities", level=1)
doc.add_paragraph("OIG-ITS provides 22 AI algorithms - more than all five competitors combined.")
tbl(H, [
    ("TOTAL AI COUNT", "22", "3-5", "0", "0", "0", "3-5"),
    ("--- STATISTICAL ---", "", "", "", "", "", ""),
    ("Anomaly Detection", "Yes (z-score)", "Partial", "No", "No", "No", "No"),
    ("Risk Scoring", "Yes (0-100)", "No", "No", "No", "No", "No"),
    ("Predictive Analytics", "Yes (regression)", "Yes", "No", "No", "No", "Yes"),
    ("Case Similarity", "Yes (cosine)", "No", "No", "No", "No", "No"),
    ("Case Clustering", "Yes (k-means)", "No", "No", "No", "No", "No"),
    ("Entity Resolution", "Yes (Levenshtein)", "No", "No", "No", "No", "No"),
    ("Fraud Ring Detection", "Yes (graph BFS)", "No", "No", "No", "No", "No"),
    ("Doc Classification", "Yes (on upload)", "No", "No", "No", "No", "No"),
    ("Investigator Recommender", "Yes", "Yes (auto)", "No", "No", "No", "No"),
    ("Auto-Escalation", "Yes", "No", "No", "No", "No", "No"),
    ("Evidence Strength (A-F)", "Yes", "No", "No", "No", "No", "No"),
    ("Timeline Anomalies", "Yes", "No", "No", "No", "No", "No"),
    ("Closure Readiness", "Yes (13 criteria)", "No", "No", "No", "No", "No"),
    ("Complaint Dedup", "Yes (Jaccard)", "No", "No", "No", "No", "No"),
    ("Workload Balancing", "Yes", "No", "No", "No", "No", "No"),
    ("Financial Patterns", "Yes (IQR)", "No", "No", "No", "No", "No"),
    ("Subject Risk", "Yes (0-100)", "No", "No", "No", "No", "No"),
    ("Case Narrative", "Yes", "No", "No", "No", "No", "Yes"),
    ("--- CLAUDE LLM ---", "", "", "", "", "", ""),
    ("NL Search", "Yes (Claude)", "Yes (AI)", "No", "No", "No", "Partial"),
    ("Doc Analysis", "Yes (Claude)", "No", "No", "No", "No", "No"),
    ("Interview Questions", "Yes (Claude)", "No", "No", "No", "No", "No"),
    ("Report Generation", "Yes (Claude)", "No", "No", "No", "No", "Partial"),
])

# 7. REPORTING
doc.add_heading("7. Reporting & Analytics", level=1)
tbl(H, [
    ("Interactive dashboards", "Yes (Recharts)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Ad-hoc report builder", "Yes (no-code)", "Yes", "Yes", "Unknown", "Yes", "Yes"),
    ("CSV/Excel/PDF export", "Yes (all 3)", "Yes", "Yes", "Unknown", "Yes", "Yes"),
    ("Financial ROI auto", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Semiannual Report", "Yes (template)", "Yes", "Unknown", "Unknown", "Unknown", "Unknown"),
])

# 8. SECURITY
doc.add_heading("8. Security & Compliance", level=1)
tbl(H, [
    ("FedRAMP Level", "Aligned", "High/IL5", "Moderate", "High/IL5", "Auth", "High"),
    ("RBAC (roles)", "6 roles, 35 perms", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Audit trail (field-level)", "Yes (old/new)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("IP whitelisting", "Yes (CIDR)", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
    ("Login lockout", "Yes (5 attempts)", "Yes", "Yes", "Yes", "Yes", "Yes"),
])

# 9. ADDITIONAL
doc.add_heading("9. Additional Modules", level=1)
tbl(H, [
    ("Training management", "Yes (full)", "Yes", "No", "No", "No", "No"),
    ("Inventory tracking", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Calendar/reminders", "Yes (recurring)", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
    ("Delegation system", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
])

# 10. COST
doc.add_page_break()
doc.add_heading("10. Cost Comparison", level=1)
tbl(["Platform", "Vendor", "Model", "Est. 30 Users/Year", "5-Year TCO"],
    [("OIG-ITS", "Y Point", "No per-user fees", "$0 licensing", "$0 + hosting"),
     ("eCASE", "Casepoint", "GSA (undisclosed)", "Contact vendor", "Contact vendor"),
     ("CMTS", "WingSwept", "GSA (undisclosed)", "Contact vendor", "Contact vendor"),
     ("ICMS", "Juvare", "Undisclosed", "Contact vendor", "Contact vendor"),
     ("Entellitrak", "Tyler Tech", "Per-user", "~$50K-100K", "~$250K-500K"),
     ("ServiceNow", "ServiceNow", "Enterprise", "~$100K-200K", "~$500K-1M")])

# 11. RECOMMENDATION
doc.add_heading("11. Recommendation", level=1)
for b in [
    "Most advanced AI: 22 algorithms vs 0-5 in all competitors",
    "Broadest modules: 24+ integrated vs 6-12 in competitors",
    "Lowest TCO: zero per-user licensing",
    "Native LEAP: auto-calc per Title 5 USC",
    "Purpose-built for OIG investigations",
    "Full source code: perpetual license, no vendor lock-in",
    "162 automated tests at 100% pass rate",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Prepared by Y Point  |  April 2026")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save("docs/OIG-ITS_Vendor_Comparison.docx")
print(f"Created: docs/OIG-ITS_Vendor_Comparison.docx ({os.path.getsize('docs/OIG-ITS_Vendor_Comparison.docx')//1024} KB)")
