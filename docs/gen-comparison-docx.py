"""Generate competitive comparison Word document with AI features."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# Cover
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OIG-ITS")
r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Investigative Tracking System")
r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x33, 0x44, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Competitive Feature Comparison & AI Capabilities Analysis")
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x25, 0x63, 0xeb); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Y Point  |  April 2026")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# 1. Executive Summary
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The OIG Investigative Tracking System (OIG-ITS) is a purpose-built investigation management "
    "platform developed for the OPM Office of Inspector General. This document compares OIG-ITS "
    "against four leading commercial platforms: Salesforce Shield/Government Cloud, Microsoft "
    "Dynamics 365, i-Sight (now Resolver), and CaseIQ (EthicsPoint)."
)
doc.add_paragraph(
    "Key findings: OIG-ITS delivers 22 AI algorithms (more than all competitors combined), "
    "eliminates per-user licensing fees, provides 24+ integrated modules versus 6-10 in competing "
    "platforms, and is purpose-built for OIG investigative workflows."
)

t = doc.add_table(rows=7, cols=2)
t.style = "Light Grid Accent 1"
for i, (k, v) in enumerate([
    ("AI Algorithms", "22 (18 statistical + 4 Claude LLM)"),
    ("Functional Modules", "24+ integrated modules"),
    ("API Routes", "120+ RESTful endpoints"),
    ("SOW Coverage", "91% (209/230 requirements)"),
    ("Automated Tests", "162/162 passing (100%)"),
    ("Per-User License Cost", "$0 (no per-user fees)"),
    ("Deployment", "On-premise Docker or Government Cloud"),
]):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v

# 2. Feature Comparison
doc.add_page_break()
doc.add_heading("2. Feature Comparison Matrix", level=1)

def add_comparison_table(title, rows):
    doc.add_heading(title, level=3)
    t = doc.add_table(rows=len(rows)+1, cols=6)
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(["Feature", "OIG-ITS", "Salesforce", "Dynamics", "i-Sight", "CaseIQ"]):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for run in p.runs: run.bold = True; run.font.size = Pt(8)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = val
            for p in t.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if j == 1: run.bold = True

add_comparison_table("Case Management", [
    ("Full case lifecycle", "Yes (7 statuses)", "Yes", "Yes", "Yes", "Yes"),
    ("Pre-close validation", "Yes (13 criteria)", "No", "No", "Partial", "No"),
    ("Case locking on close", "Yes (auto)", "Partial", "Partial", "Yes", "Partial"),
    ("Field permissions by status", "Yes (Role x Status)", "Partial", "Partial", "No", "No"),
    ("Case relationships", "Yes", "Yes", "Yes", "Yes", "Partial"),
    ("Draft / autosave", "Yes", "Yes", "Yes", "No", "No"),
])

add_comparison_table("Evidence & Documents", [
    ("Immutable chain of custody", "Yes", "No", "No", "Yes", "Partial"),
    ("Exhibit labeling (EX-001)", "Yes", "No", "No", "Yes", "No"),
    ("Document versioning", "Yes (v1,v2)", "Yes", "Yes", "Partial", "Partial"),
    ("19 document templates", "Yes", "Partial", "Partial", "Partial", "Partial"),
    ("ZIP download", "Yes", "No", "No", "Yes", "No"),
])

add_comparison_table("Public Intake", [
    ("Public hotline (no auth)", "Yes", "Partial", "No", "Yes", "Yes"),
    ("Whistleblower protections", "Yes (WPA cited)", "No", "No", "Partial", "Yes"),
    ("AI risk scoring on intake", "Yes (0-100)", "No", "No", "No", "No"),
    ("Inquiry to Case (1-click)", "Yes", "Partial", "Partial", "Yes", "Yes"),
])

add_comparison_table("Reporting & Analytics", [
    ("Interactive dashboards", "Yes (Recharts)", "Yes", "Yes", "Yes", "Yes"),
    ("Ad-hoc report builder", "Yes (no-code)", "Yes", "Yes", "Yes", "Partial"),
    ("Financial ROI auto-calc", "Yes", "Partial", "Partial", "Partial", "No"),
    ("Congressional report template", "Yes", "No", "No", "No", "No"),
])

add_comparison_table("Additional Modules", [
    ("Training management", "Yes (full)", "No", "No", "No", "No"),
    ("Time & labor tracking", "Yes + avail pay", "No", "No", "No", "No"),
    ("Inventory tracking", "Yes", "No", "No", "No", "No"),
])

# 3. AI Detailed Comparison
doc.add_page_break()
doc.add_heading("3. AI/ML Capabilities - Detailed Comparison", level=1)
doc.add_paragraph(
    "AI capabilities are the primary differentiator. OIG-ITS provides 22 algorithms - "
    "more than all competitors combined."
)

ai_rows = [
    ("AI Algorithm Count", "22", "3-5", "2-3", "0", "0"),
    ("Anomaly Detection", "Yes (z-score)", "Partial", "No", "No", "No"),
    ("Risk Scoring", "Yes (0-100)", "No", "No", "No", "No"),
    ("Predictive Analytics", "Yes (regression)", "Partial", "Partial", "No", "No"),
    ("Case Similarity", "Yes (cosine)", "Partial", "No", "No", "No"),
    ("Case Clustering", "Yes (k-means)", "No", "No", "No", "No"),
    ("Entity Resolution", "Yes (Levenshtein)", "No", "No", "No", "No"),
    ("Fraud Ring Detection", "Yes (graph BFS)", "No", "No", "No", "No"),
    ("Document Classification", "Yes (on upload)", "Partial", "No", "No", "No"),
    ("Investigator Recommender", "Yes", "No", "No", "No", "No"),
    ("Auto-Escalation", "Yes", "No", "No", "No", "No"),
    ("Evidence Strength (A-F)", "Yes", "No", "No", "No", "No"),
    ("Timeline Anomalies", "Yes", "No", "No", "No", "No"),
    ("Closure Readiness", "Yes (13 criteria)", "No", "No", "No", "No"),
    ("Complaint Dedup", "Yes (Jaccard)", "No", "No", "No", "No"),
    ("Workload Balancing", "Yes", "No", "No", "No", "No"),
    ("Financial Patterns", "Yes (IQR)", "No", "No", "No", "No"),
    ("Subject Risk (0-100)", "Yes", "No", "No", "No", "No"),
    ("Case Narrative", "Yes", "No", "No", "No", "No"),
    ("NL Search (LLM)", "Yes (Claude)", "Partial", "Partial", "No", "No"),
    ("Doc Analysis (LLM)", "Yes (Claude)", "No", "No", "No", "No"),
    ("Interview Questions (LLM)", "Yes (Claude)", "No", "No", "No", "No"),
    ("Report Generation (LLM)", "Yes (Claude)", "No", "Partial", "No", "No"),
]
add_comparison_table("All 22 AI Algorithms", ai_rows)

# 4. AI Inventory
doc.add_page_break()
doc.add_heading("4. OIG-ITS AI Algorithm Inventory", level=1)

doc.add_heading("Statistical & Rule-Based (18)", level=2)
stat = [
    ("1. Anomaly Detection", "Z-score", "Financial outliers, long cases, abnormal activity"),
    ("2. Risk Scoring", "Keyword+rules", "Score 0-100 on complaints, auto-escalate priority"),
    ("3. Predictive Analytics", "Regression", "Case duration, outcome, caseload forecast"),
    ("4. Case Similarity", "Cosine sim", "22-dim vectors, top N similar cases"),
    ("5. Case Clustering", "K-means k=5", "Group cases by feature similarity"),
    ("6. Entity Resolution", "Levenshtein+Soundex", "Fuzzy name match, phonetic, email domain"),
    ("7. Network Analysis", "Graph BFS", "Hubs, components, fraud rings"),
    ("8. Doc Classifier", "Regex+MIME", "8 categories, auto-tags, runs on upload"),
    ("9. Investigator Recommender", "Multi-factor", "Expertise 30pt, workload 25pt, success 25pt"),
    ("10. Auto-Escalation", "Rules", "Financial >$500K, deadline <7d, stale 14d"),
    ("11. Evidence Strength", "Weighted", "0-100, A-F grade, 5 scoring factors"),
    ("12. Timeline Anomalies", "Gap analysis", "Gaps >30d, post-review evidence, weekend"),
    ("13. Closure Readiness", "Checklist", "13 criteria, 100 points, ready at 80+"),
    ("14. Complaint Dedup", "Jaccard+Levenshtein", "60% subject + 40% keyword overlap"),
    ("15. Workload Balancing", "Std dev", "Overloaded >avg+1.5sigma, queue depth"),
    ("16. Financial Patterns", "IQR mining", "Round numbers, thresholds, outliers"),
    ("17. Subject Risk", "Multi-factor", "Cases, violations, financial, network, repeat"),
    ("18. Case Narrative", "Template", "Auto-generate investigation summary"),
]

t = doc.add_table(rows=len(stat)+1, cols=3)
t.style = "Light Grid Accent 1"
for j, h in enumerate(["Algorithm", "Method", "Description"]):
    t.rows[0].cells[j].text = h
    for p in t.rows[0].cells[j].paragraphs:
        for run in p.runs: run.bold = True; run.font.size = Pt(8)
for i, (name, method, desc) in enumerate(stat):
    t.rows[i+1].cells[0].text = name
    t.rows[i+1].cells[1].text = method
    t.rows[i+1].cells[2].text = desc
    for c in t.rows[i+1].cells:
        for p in c.paragraphs:
            for run in p.runs: run.font.size = Pt(8)

doc.add_heading("Claude LLM-Powered (4)", level=2)
for name, desc in [
    ("19. Natural Language Search", "Plain English to structured filters via Claude Sonnet 4"),
    ("20. Document Analysis", "Extract facts, entities, red flags from document text"),
    ("21. Interview Questions", "Role-appropriate questions: opening, substantive, probe, closing"),
    ("22. Report Generation", "Professional reports (summary/narrative/findings) from case data"),
]:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

# 5. Cost Analysis
doc.add_page_break()
doc.add_heading("5. Cost Analysis (30 Users, 5 Years)", level=1)
t = doc.add_table(rows=6, cols=4)
t.style = "Light Grid Accent 1"
for i, row in enumerate([
    ["Platform", "Per User/Mo", "30 Users/Year", "5-Year Total"],
    ["OIG-ITS", "$0", "$0", "$0 licensing"],
    ["Salesforce", "$150-300", "$54K-108K", "$270K-540K"],
    ["Dynamics 365", "$95-210", "$34K-76K", "$170K-378K"],
    ["i-Sight", "$50-100", "$18K-36K", "$90K-180K"],
    ["CaseIQ", "Custom", "~$30K-60K", "~$150K-300K"],
]):
    for j, v in enumerate(row):
        t.rows[i].cells[j].text = v
        for p in t.rows[i].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                if i == 0: run.bold = True
                if i == 1: run.bold = True; run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

# 6. Recommendation
doc.add_heading("6. Recommendation", level=1)
for b in [
    "Highest AI capabilities: 22 algorithms vs 0-5 in competitors",
    "Lowest cost: zero per-user licensing ($0 vs $90K-540K over 5 years)",
    "Most integrated: 24+ modules vs 6-10 in competitors",
    "Purpose-built for OIG: not adapted from generic CRM",
    "Full source code: no vendor lock-in",
    "91% SOW coverage, 162 tests passing at 100%",
]:
    doc.add_paragraph(b, style="List Bullet")

p = doc.add_paragraph()
r = p.add_run("Prepared by Y Point  |  April 2026")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save("docs/OIG-ITS_Competitive_Comparison.docx")
import os
print(f"Created: docs/OIG-ITS_Competitive_Comparison.docx ({os.path.getsize('docs/OIG-ITS_Competitive_Comparison.docx')//1024} KB)")
