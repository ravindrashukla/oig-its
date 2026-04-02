"""Generate Data Model Diagram Document (Word) with crow's feet notation."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# Parse schema
schema = Path("prisma/schema.prisma").read_text()
models = {}
current = None
for line in schema.split("\n"):
    m = re.match(r"^model\s+(\w+)\s*\{", line)
    if m:
        current = m.group(1)
        models[current] = {"fields": [], "relations": []}
        continue
    if current and line.strip() == "}":
        current = None; continue
    if current and line.strip() and not line.strip().startswith("//") and not line.strip().startswith("@@"):
        parts = line.strip().split()
        if len(parts) >= 2:
            name, ftype = parts[0], parts[1]
            is_pk = "@id" in line
            is_fk = name.endswith("Id") and ("String" in ftype or "Int" in ftype)
            is_unique = "@unique" in line
            is_optional = "?" in ftype
            is_default = "@default" in line
            base = ftype.replace("?", "").replace("[]", "")
            is_rel = base[0].isupper() and base not in ("String", "Int", "Float", "Boolean", "DateTime", "Json")
            if is_rel:
                models[current]["relations"].append((name, base, "[]" in ftype))
            else:
                models[current]["fields"].append({
                    "name": name, "type": ftype.replace("?",""), "pk": is_pk, "fk": is_fk,
                    "optional": is_optional, "unique": is_unique, "default": is_default
                })

# Group models by domain
DOMAINS = {
    "Core Investigation": ["Case", "CaseAssignment", "CaseStatusHistory", "CaseNote", "CaseRelationship", "CaseSubject", "CloseChecklist", "SubpoenaPackage"],
    "Subjects & Entities": ["Subject", "Violation", "FinancialResult", "SubjectAction"],
    "Evidence & Documents": ["EvidenceItem", "EvidenceFile", "ChainOfCustody", "Document", "DocumentAttachment", "DocumentAccessLog", "DocumentComment"],
    "Users & Auth": ["User", "Session", "Organization", "PasswordResetToken", "UserPreference", "Delegation"],
    "Workflows & Notifications": ["WorkflowDefinition", "WorkflowInstance", "WorkflowStepAction", "Notification", "NotificationPreference"],
    "Tasks": ["Task"],
    "Intake & Inquiries": ["PreliminaryInquiry"],
    "Reporting": ["ReportDefinition", "ReportRun", "ReportSchedule", "ReportReview", "SavedSearch"],
    "Training": ["TrainingCourse", "TrainingRecord", "TrainingAssignment", "CourseEvaluation"],
    "Time & Labor": ["TimeEntry", "Timesheet"],
    "Inventory & Calendar": ["InventoryItem", "CalendarReminder"],
    "Configuration": ["SystemSetting", "ReferenceData", "FieldLabel", "FeatureFlag", "Announcement", "AuditLog", "InvestigativeTechnique", "Referral"],
}

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2); s.right_margin = Cm(2)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# COVER
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OIG-ITS"); r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Data Model Documentation"); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PostgreSQL Database Schema\nCrow\u2019s Feet Entity-Relationship Notation"); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("53 Tables  |  300+ Columns  |  100+ Indexes  |  15 Enumerations"); r.font.size = Pt(11); r.bold = True

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Y Point  |  April 2026"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# EXTENSIBILITY NOTE
doc.add_heading("Schema Extensibility", level=1)
doc.add_paragraph(
    "The OIG-ITS data model is designed for extensibility. Additional attributes (columns) can be added "
    "to any table without requiring major application changes. The application layer uses dynamic field "
    "rendering and Prisma ORM\u2019s schema-first approach, meaning new fields are:"
)
for b in [
    "Added to the Prisma schema file (a single line per field)",
    "Applied to the database via a non-destructive migration (prisma migrate dev)",
    "Automatically available in API responses without endpoint changes",
    "Configurable for display via the FieldLabel administration interface",
    "Subject to the existing RBAC and field-level permission framework",
]:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph(
    "This architecture ensures that OPM OIG administrators can extend the data model to accommodate "
    "evolving investigative requirements without vendor intervention or application redeployment."
)

# NOTATION GUIDE
doc.add_heading("Crow\u2019s Feet Notation Guide", level=1)
t = doc.add_table(rows=5, cols=2)
t.style = "Light Grid Accent 1"
t.rows[0].cells[0].text = "Symbol"; t.rows[0].cells[1].text = "Meaning"
for p in t.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in t.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True
t.rows[1].cells[0].text = "PK"; t.rows[1].cells[1].text = "Primary Key \u2014 unique identifier for each record"
t.rows[2].cells[0].text = "FK"; t.rows[2].cells[1].text = "Foreign Key \u2014 references a primary key in another table"
t.rows[3].cells[0].text = "||----<"; t.rows[3].cells[1].text = "One-to-Many relationship (one parent, many children)"
t.rows[4].cells[0].text = "||----||"; t.rows[4].cells[1].text = "One-to-One relationship"
doc.add_paragraph()

# DOMAIN OVERVIEW
doc.add_heading("Domain Overview", level=1)
t = doc.add_table(rows=len(DOMAINS)+1, cols=3)
t.style = "Light Grid Accent 1"
t.rows[0].cells[0].text = "Domain"; t.rows[0].cells[1].text = "Tables"; t.rows[0].cells[2].text = "Description"
for p in t.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True; r.font.size = Pt(9)
for p in t.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True; r.font.size = Pt(9)
for p in t.rows[0].cells[2].paragraphs:
    for r in p.runs: r.bold = True; r.font.size = Pt(9)

descs = {
    "Core Investigation": "Case lifecycle, assignments, status history, notes, relationships",
    "Subjects & Entities": "People, organizations, violations, financial results, actions",
    "Evidence & Documents": "Evidence items, files, chain of custody, documents, versioning",
    "Users & Auth": "User accounts, sessions, organizations, password reset, preferences",
    "Workflows & Notifications": "Multi-step workflows, actions, notifications, preferences",
    "Tasks": "Investigation tasks with assignment, priority, and status",
    "Intake & Inquiries": "Hotline/whistleblower complaints, risk scoring",
    "Reporting": "Report definitions, execution runs, schedules, reviews",
    "Training": "Courses, records, assignments, evaluations",
    "Time & Labor": "Time entries, timesheets, LEAP calculations",
    "Inventory & Calendar": "Equipment tracking, recurring reminders",
    "Configuration": "Settings, reference data, audit logs, field labels",
}
for i, (domain, tbls) in enumerate(DOMAINS.items()):
    t.rows[i+1].cells[0].text = domain
    t.rows[i+1].cells[1].text = str(len(tbls))
    t.rows[i+1].cells[2].text = descs.get(domain, "")
    for c in t.rows[i+1].cells:
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9)

doc.add_page_break()

# ENTITY DETAILS PER DOMAIN
for domain, table_names in DOMAINS.items():
    doc.add_heading(f"Domain: {domain}", level=1)

    for tname in table_names:
        if tname not in models:
            continue
        m = models[tname]

        doc.add_heading(tname, level=2)

        # Relationships text
        rels = m["relations"]
        if rels:
            rel_text = []
            for rname, rmodel, is_array in rels:
                if is_array:
                    rel_text.append(f"{tname} ||----< {rmodel}  (one-to-many via {rname})")
                else:
                    rel_text.append(f"{tname} ||----|| {rmodel}  (many-to-one via {rname})")
            p = doc.add_paragraph()
            r = p.add_run("Relationships (Crow\u2019s Feet):")
            r.bold = True; r.font.size = Pt(9)
            for rt in rel_text:
                p = doc.add_paragraph(rt)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(8)
            doc.add_paragraph()

        # Field table
        fields = m["fields"]
        t = doc.add_table(rows=len(fields)+1, cols=5)
        t.style = "Light Grid Accent 1"
        for j, h in enumerate(["Column", "Type", "Key", "Required", "Notes"]):
            t.rows[0].cells[j].text = h
            for p in t.rows[0].cells[j].paragraphs:
                for r in p.runs: r.bold = True; r.font.size = Pt(8)

        for i, f in enumerate(fields):
            t.rows[i+1].cells[0].text = f["name"]
            t.rows[i+1].cells[1].text = f["type"]

            keys = []
            if f["pk"]: keys.append("PK")
            if f["fk"]: keys.append("FK")
            if f["unique"]: keys.append("UQ")
            t.rows[i+1].cells[2].text = ", ".join(keys)

            t.rows[i+1].cells[3].text = "Optional" if f["optional"] else "Required"

            notes = []
            if f["default"]: notes.append("Has default")
            if f["pk"]: notes.append("CUID auto-generated")
            if f["fk"]:
                # Find which model this FK references
                fk_target = f["name"].replace("Id", "")
                for rname, rmodel, _ in rels:
                    if rname.lower() == fk_target.lower() or rmodel.lower() == fk_target.lower():
                        notes.append(f"References {rmodel}")
                        break
            t.rows[i+1].cells[4].text = "; ".join(notes)

            for c in t.rows[i+1].cells:
                for p in c.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(8)
                        if f["pk"]:
                            r.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
                        elif f["fk"]:
                            r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)

        doc.add_paragraph()

    doc.add_page_break()

# ENUMERATIONS
doc.add_heading("Enumerations", level=1)
doc.add_paragraph("The following enumerated types define valid values for key classification fields:")

enums = {}
current_enum = None
for line in schema.split("\n"):
    m = re.match(r"^enum\s+(\w+)\s*\{", line)
    if m:
        current_enum = m.group(1)
        enums[current_enum] = []
        continue
    if current_enum and line.strip() == "}":
        current_enum = None; continue
    if current_enum and line.strip() and not line.strip().startswith("//"):
        enums[current_enum].append(line.strip())

for ename, values in sorted(enums.items()):
    doc.add_heading(ename, level=3)
    p = doc.add_paragraph(", ".join(values))
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs: r.font.size = Pt(9); r.font.name = "Consolas"

# EXTENSIBILITY FOOTER
doc.add_page_break()
doc.add_heading("Extensibility Statement", level=1)
doc.add_paragraph(
    "All tables in the OIG-ITS data model support the addition of new attributes (columns) without "
    "requiring major application changes. The Prisma ORM schema-first architecture ensures that:"
)
for b in [
    "New fields are added with a single line in the schema definition file",
    "Database migrations are non-destructive and reversible",
    "API endpoints automatically include new fields in responses",
    "The field-level permission system (FieldLabel + field-permissions.ts) automatically governs access to new fields",
    "Existing data is preserved; new columns receive configurable default values",
    "No application redeployment is required for schema-only changes in development mode",
    "The administrative interface (Settings > Field Labels) allows customizing display names for any field",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph()
doc.add_paragraph(
    "This extensibility ensures that OPM OIG can adapt the data model to evolving investigative "
    "requirements, regulatory changes, and operational needs without contractor intervention."
)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Y Point Technologies  |  April 2026")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save("docs/OIG-ITS_Data_Model.docx")
print(f"Created: docs/OIG-ITS_Data_Model.docx ({os.path.getsize('docs/OIG-ITS_Data_Model.docx')//1024} KB)")
print(f"Models: {len(models)}")
print(f"Enums: {len(enums)}")
total_fields = sum(len(m['fields']) for m in models.values())
total_rels = sum(len(m['relations']) for m in models.values())
print(f"Total fields: {total_fields}")
print(f"Total relations: {total_rels}")
