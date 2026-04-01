"""Convert all markdown docs to Word (.docx) format."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = Path(__file__).parent

def md_to_docx(md_path, docx_path, title):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    in_table = False
    table_rows = []
    in_code = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = doc.styles['Normal']
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Table handling
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Check if separator row
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue

            if not in_table:
                in_table = True
                table_rows = []

            table_rows.append(cells)

            # Check if next line is NOT a table row
            next_is_table = (i + 1 < len(lines) and '|' in lines[i+1] and lines[i+1].strip().startswith('|'))
            if not next_is_table:
                # Render the table
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                    tbl.style = 'Light Grid Accent 1'
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = tbl.rows[row_idx].cells[col_idx]
                                cell.text = cell_text.replace('**', '').replace('`', '')
                                for paragraph in cell.paragraphs:
                                    paragraph.paragraph_format.space_before = Pt(1)
                                    paragraph.paragraph_format.space_after = Pt(1)
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)

                    doc.add_paragraph()  # spacing after table

                in_table = False
                table_rows = []

            i += 1
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # remove bold markers
            text = re.sub(r'`(.*?)`', r'\1', text)  # remove code markers
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if m:
            text = m.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # remove links

        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(4)

        i += 1

    doc.save(str(docx_path))
    print(f"  {docx_path.name} ({docx_path.stat().st_size // 1024} KB)")


# Convert all three documents
files = [
    ("technical-specification.md", "OIG-ITS_Technical_Specification.docx", "Technical Specification"),
    ("user-manual.md", "OIG-ITS_User_Manual.docx", "User Manual"),
    ("implementation-plan.md", "OIG-ITS_Implementation_Plan.docx", "Implementation Plan"),
]

print("Converting markdown to Word...")
for md_name, docx_name, title in files:
    md_path = DOCS_DIR / md_name
    docx_path = DOCS_DIR / docx_name
    if md_path.exists():
        md_to_docx(md_path, docx_path, title)
    else:
        print(f"  SKIP: {md_name} not found")

print("Done!")
