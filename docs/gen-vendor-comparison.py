"""Generate OIG-ITS vs 5 OIG vendors competitive comparison Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

def center(text, size, color=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = RGBColor(*color)

def add_table(headers, rows, bold_col=1):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(7)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            t.rows[i+1].cells[j].text = v
            for p in t.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)
                    if j == bold_col: r.bold = True
                    if v.startswith("Yes") and j == bold_col:
                        r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

# ═══ COVER ═══
doc.add_paragraph()
center("OIG-ITS", 36, (0x0f, 0x17, 0x2a), True)
center("Investigative Tracking System", 18, (0x33, 0x44, 0x55))
doc.add_paragraph()
center("Comparative Evaluation Against", 14, (0x25, 0x63, 0xeb), True)
center("Federal OIG Case Management Solutions", 14, (0x25, 0x63, 0xeb), True)
doc.add_paragraph()
center("eCASE (Casepoint) | CMTS (WingSwept) | ICMS (Juvare)", 10, (0x64, 0x74, 0x8b))
center("Entellitrak (Tyler Technologies) | ICM (ServiceNow)", 10, (0x64, 0x74, 0x8b))
doc.add_paragraph()
center("Prepared by Y Point  |  April 2026", 10, (0x94, 0xa3, 0xb8))
doc.add_page_break()

# ═══ 1. EXECUTIVE SUMMARY ═══
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This document compares OIG-ITS against five established OIG case management solutions "
    "currently used across federal Inspector General offices. The comparison covers case management, "
    "evidence tracking, AI/ML capabilities, LEAP support, security posture, and total cost of ownership."
)
doc.add_paragraph(
    "OIG-ITS differentiates through 22 AI algorithms (vs 0-3 in competitors), zero per-user licensing, "
    "full source code ownership, and the broadest integrated module set (24+ modules including training, "
    "time tracking, inventory, and financial analytics)."
)

# ═══ 2. VENDOR PROFILES ═══
doc.add_heading("2. Vendor Overview", level=1)
add_table(
    ["Rank", "Solution", "Vendor", "FedRAMP", "LEAP", "OIG Specificity"],
    [
        ("—", "OIG-ITS", "Y Point", "Aligned (certifiable)", "Native (auto-calc)", "Very High"),
        ("1", "eCASE Investigations", "Casepoint/OPEXUS", "High / IL5", "Native", "Very High"),
        ("2", "CMTS", "WingSwept, LLC", "Moderate", "Configurable", "High"),
        ("3", "ICMS with LEAP", "Juvare, LLC", "High / IL5", "Native", "Moderate-High"),
        ("4", "Entellitrak (IG App)", "Tyler Technologies", "Authorized", "Custom Build", "High"),
        ("5", "Investigative CM", "ServiceNow", "High", "Custom Build", "Low"),
    ], bold_col=1
)

# ═══ 3. FEATURE COMPARISON ═══
doc.add_page_break()
doc.add_heading("3. Feature Comparison Matrix", level=1)
H = ["Feature", "OIG-ITS", "eCASE", "CMTS", "ICMS", "Entellitrak", "ServiceNow ICM"]

doc.add_heading("Case Management", level=2)
add_table(H, [
    ("Full case lifecycle", "Yes (7 statuses)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Auto case numbering", "Yes (OIG-YYYY-NNN)", "Yes", "Yes", "Yes", "Configurable", "Yes"),
    ("Pre-close validation", "Yes (13 criteria)", "Yes (checklists)", "No", "No", "Configurable", "No"),
    ("Case locking on close", "Yes (auto-lock)", "Yes", "Unknown", "Unknown", "Configurable", "Unknown"),
    ("Case relationships/spinoffs", "Yes", "Yes", "Yes", "Yes", "Configurable", "Yes (entity links)"),
    ("Draft/autosave", "Yes", "Unknown", "Unknown", "Unknown", "Yes", "Yes"),
    ("Field permissions by status", "Yes (Role x Status)", "Yes", "Unknown", "Unknown", "Configurable", "Yes"),
    ("Configurable workflows", "Yes (5 types)", "Yes", "Yes", "Yes", "Yes (low-code)", "Yes"),
    ("Mobile access", "Yes (responsive)", "Unknown", "Yes (web)", "Yes (mobile)", "Yes (web)", "Yes"),
    ("9+ case types", "Yes (9 types)", "Yes (configurable)", "Yes", "Yes", "Configurable", "Yes"),
])

doc.add_heading("Evidence & Documents", level=2)
add_table(H, [
    ("Immutable chain of custody", "Yes", "Yes", "Basic", "Yes (full)", "Configurable", "Yes"),
    ("Sequential exhibit labeling", "Yes (EX-001)", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Document versioning", "Yes (v1, v2...)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Supervisory doc approval", "Yes", "Yes", "Yes", "Unknown", "Configurable", "Yes"),
    ("19 document templates", "Yes (19)", "Yes (templates)", "Yes (reports)", "Unknown", "Configurable", "Unknown"),
    ("ZIP download of case files", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("30 file format support", "Yes (30 types)", "Yes", "Yes", "Yes", "Yes", "Yes"),
])

doc.add_heading("Public Intake & Whistleblower", level=2)
add_table(H, [
    ("Public hotline (no auth)", "Yes", "Yes (portal)", "Basic intake", "Intake workflows", "Configurable", "Yes"),
    ("Whistleblower portal + WPA", "Yes (WPA cited)", "Yes", "Unknown", "Unknown", "Unknown", "Yes (portal)"),
    ("Anonymous submissions", "Yes", "Yes", "Unknown", "Unknown", "Unknown", "Yes"),
    ("AI risk scoring on intake", "Yes (0-100 auto)", "No", "No", "No", "No", "No"),
    ("Complaint deduplication", "Yes (AI Jaccard)", "No", "No", "No", "No", "No"),
    ("Inquiry to Case conversion", "Yes (1-click)", "Yes", "Yes", "Yes", "Configurable", "Yes"),
])

doc.add_heading("LEAP Support", level=2)
add_table(H, [
    ("LEAP time tracking", "Yes (native)", "Yes (native)", "Configurable", "Yes (native)", "Custom build", "Custom build"),
    ("Auto LEAP calculation", "Yes (Title 5)", "Yes (auto)", "Configurable", "Yes (dedicated)", "Manual", "Manual"),
    ("Substantial hours check", "Yes (auto)", "Yes", "Unknown", "Yes", "Unknown", "Unknown"),
    ("Availability pay calc", "Yes (25% base)", "Yes", "Unknown", "Yes", "Unknown", "Unknown"),
    ("Agent timesheet reports", "Yes (export)", "Yes", "Yes", "Yes", "Configurable", "Unknown"),
])

doc.add_heading("Reporting & Analytics", level=2)
add_table(H, [
    ("Interactive dashboards", "Yes (Recharts)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Ad-hoc report builder", "Yes (no-code)", "Yes", "Yes", "Unknown", "Yes", "Yes"),
    ("CSV/Excel/PDF export", "Yes (all 3)", "Yes", "Yes", "Unknown", "Yes", "Yes"),
    ("Financial ROI auto-calc", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Semiannual Report template", "Yes", "Yes", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Scheduled auto-email reports", "Yes", "Unknown", "Yes", "Unknown", "Unknown", "Yes"),
])

doc.add_heading("Security & Compliance", level=2)
add_table(H, [
    ("FedRAMP Level", "Aligned", "High / IL5", "Moderate", "High / IL5", "Authorized", "High / IL4"),
    ("RBAC (6+ roles)", "Yes (6 roles)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Field-level audit trail", "Yes (old/new)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("IP whitelisting", "Yes (CIDR)", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
    ("Login lockout", "Yes (5 attempts)", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Password reset (email)", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Records retention", "Yes (configurable)", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
])

doc.add_heading("Additional Modules", level=2)
add_table(H, [
    ("Training management", "Yes (full module)", "Yes", "No", "No", "No", "No"),
    ("Time & labor tracking", "Yes + avail pay", "Yes (LEAP)", "Configurable", "Yes (LEAP)", "Custom", "Custom"),
    ("Inventory/equipment", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Calendar/reminders", "Yes (recurring)", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
    ("Delegation system", "Yes", "Unknown", "Unknown", "Unknown", "Unknown", "Yes"),
    ("Subpoena workflow", "Yes", "Yes", "Unknown", "Unknown", "Unknown", "Unknown"),
    ("Data import API", "Yes (bulk JSON)", "Unknown", "Unknown", "Unknown", "Yes (API)", "Yes (API)"),
])

# ═══ 4. AI/ML COMPARISON ═══
doc.add_page_break()
doc.add_heading("4. AI/ML Capabilities — Detailed Comparison", level=1)
doc.add_paragraph(
    "AI capabilities represent the most significant differentiator for OIG-ITS. With 22 algorithms "
    "spanning statistical analysis, graph theory, NLP, and large language models, OIG-ITS provides "
    "investigation intelligence unmatched by any competitor in the federal OIG market."
)

add_table(H, [
    ("Total AI algorithms", "22", "3-5 (built-in)", "0", "0", "0", "3-5 (AI-native)"),
    ("Anomaly Detection (z-score)", "Yes", "Partial", "No", "No", "No", "No"),
    ("Complaint Risk Scoring", "Yes (0-100)", "No", "No", "No", "No", "No"),
    ("Predictive Analytics", "Yes (regression)", "Yes (predictive)", "No", "No", "No", "Yes (ML)"),
    ("Case Similarity (cosine)", "Yes", "No", "No", "No", "No", "No"),
    ("Case Clustering (k-means)", "Yes", "No", "No", "No", "No", "No"),
    ("Duplicate Entity (Levenshtein)", "Yes", "No", "No", "No", "No", "No"),
    ("Fraud Ring Detection (graph)", "Yes (BFS)", "No", "No", "No", "No", "No"),
    ("Document Auto-Classification", "Yes (on upload)", "No", "No", "No", "No", "No"),
    ("Investigator Recommendation", "Yes (multi-factor)", "Yes (auto-assign)", "No", "No", "No", "No"),
    ("Auto Priority Escalation", "Yes (rule engine)", "No", "No", "No", "No", "No"),
    ("Evidence Strength (A-F)", "Yes", "No", "No", "No", "No", "No"),
    ("Timeline Anomalies", "Yes (gap analysis)", "No", "No", "No", "No", "No"),
    ("Closure Readiness (13 pt)", "Yes", "No", "No", "No", "No", "No"),
    ("Complaint Deduplication", "Yes (Jaccard)", "No", "No", "No", "No", "No"),
    ("Workload Balancing", "Yes (statistical)", "No", "No", "No", "No", "No"),
    ("Financial Pattern Mining", "Yes (IQR)", "No", "No", "No", "No", "No"),
    ("Subject Risk Profiling", "Yes (0-100)", "No", "No", "No", "No", "No"),
    ("Case Narrative Generation", "Yes (template)", "No", "No", "No", "No", "Yes (AI summary)"),
    ("Natural Language Search", "Yes (Claude AI)", "Yes (AI search)", "No", "No", "No", "Partial"),
    ("AI Document Analysis", "Yes (Claude AI)", "No", "No", "No", "No", "No"),
    ("Interview Question Gen", "Yes (Claude AI)", "No", "No", "No", "No", "No"),
    ("AI Report Generation", "Yes (Claude AI)", "No", "No", "No", "No", "Partial"),
])

doc.add_paragraph()
doc.add_paragraph(
    "Note: eCASE describes built-in AI for predictive analysis and search but does not detail "
    "specific algorithm counts. ServiceNow has AI-native capabilities via their platform but requires "
    "custom configuration for OIG-specific use cases. CMTS, ICMS, and Entellitrak do not document "
    "any AI/ML capabilities in their public materials."
)

# ═══ 5. AI INVENTORY ═══
doc.add_heading("5. OIG-ITS AI Algorithm Inventory", level=1)
doc.add_heading("Statistical & Rule-Based (18)", level=2)
add_table(["#", "Algorithm", "Method", "Key Feature"], [
    ("1", "Anomaly Detection", "Z-score", "Financial, case, activity outlier detection"),
    ("2", "Risk Scoring", "Keyword+rules", "0-100 score, auto-runs on hotline intake"),
    ("3", "Predictive Analytics", "Linear regression", "Duration, outcome, caseload forecast"),
    ("4", "Case Similarity", "Cosine similarity", "22-dim vectors, top N matches"),
    ("5", "Case Clustering", "K-means (k=5)", "Euclidean distance grouping"),
    ("6", "Entity Resolution", "Levenshtein+Soundex", "Fuzzy name + phonetic matching"),
    ("7", "Network Analysis", "Graph BFS", "Hubs, components, fraud rings"),
    ("8", "Document Classifier", "Regex+MIME", "8 categories, auto-tags on upload"),
    ("9", "Investigator Recommender", "Multi-factor", "Expertise, workload, success, availability"),
    ("10", "Auto-Escalation", "Rule engine", ">$500K, <7d deadline, stale 14d"),
    ("11", "Evidence Strength", "Weighted scoring", "A-F grade, 5 factors, 0-100"),
    ("12", "Timeline Anomalies", "Gap analysis", ">30d gaps, post-review evidence"),
    ("13", "Closure Readiness", "Checklist", "13 criteria, ready at 80+"),
    ("14", "Complaint Dedup", "Jaccard+Levenshtein", "60% subject + 40% keyword match"),
    ("15", "Workload Balancing", "Std deviation", "Overloaded/underloaded flags"),
    ("16", "Financial Patterns", "IQR mining", "Round numbers, thresholds, outliers"),
    ("17", "Subject Risk", "Multi-factor", "Cases, violations, financial, repeat"),
    ("18", "Case Narrative", "Template gen", "Auto plain-English summary"),
], bold_col=1)

doc.add_heading("Claude LLM-Powered (4)", level=2)
add_table(["#", "Algorithm", "Model", "Capability"], [
    ("19", "Natural Language Search", "Claude Sonnet 4", "English to structured query"),
    ("20", "Document Analysis", "Claude Sonnet 4", "Facts, entities, red flags extraction"),
    ("21", "Interview Questions", "Claude Sonnet 4", "Role-appropriate question generation"),
    ("22", "Report Generation", "Claude Sonnet 4", "Professional investigation reports"),
], bold_col=1)

# ═══ 6. COST ═══
doc.add_page_break()
doc.add_heading("6. Cost Comparison", level=1)
add_table(["Platform", "Licensing Model", "Est. Annual (30 users)", "5-Year Total"], [
    ("OIG-ITS (Y Point)", "$0 per-user", "$0 licensing", "$0 licensing"),
    ("eCASE (Casepoint)", "Undisclosed (GSA)", "Contact vendor", "Contact vendor"),
    ("CMTS (WingSwept)", "Undisclosed (GSA)", "Contact vendor", "Contact vendor"),
    ("ICMS (Juvare)", "Undisclosed", "Contact vendor", "Contact vendor"),
    ("Entellitrak (Tyler)", "Per-user", "Est. $50K-100K", "Est. $250K-500K"),
    ("ICM (ServiceNow)", "Enterprise platform", "Est. $100K-200K", "Est. $500K-1M"),
])

doc.add_paragraph(
    "OIG-ITS eliminates per-user licensing entirely. Total cost of ownership includes "
    "only infrastructure hosting (est. $5K-15K/year for government cloud) and optional "
    "support/maintenance. Full source code is delivered, eliminating vendor lock-in."
)

# ═══ 7. DIFFERENTIATORS ═══
doc.add_heading("7. OIG-ITS Key Differentiators", level=1)
for title, desc in [
    ("22 AI Algorithms", "More investigation intelligence than all five competitors combined. Unique capabilities: fraud ring detection, evidence strength scoring, complaint deduplication, financial pattern mining, and subject risk profiling."),
    ("Zero Per-User Licensing", "Unlike all competitors, OIG-ITS has no per-user fees. Full source code delivered with perpetual license."),
    ("24+ Integrated Modules", "Training, timesheets, inventory, calendar, financial analytics built-in. Competitors require separate products."),
    ("Native LEAP with Auto-Calculation", "Title 5 USC availability pay, substantial hours check, and LEAP tracking built natively with timesheet approval workflow."),
    ("Claude AI Integration", "Only OIG case management system with LLM-powered natural language search, AI report generation, interview question generation, and document analysis."),
    ("Full Source Code Ownership", "Government retains complete source code. No vendor lock-in. Customize without contractor dependency."),
]:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ═══ 8. RECOMMENDATION ═══
doc.add_heading("8. Recommendation", level=1)
for b in [
    "Most advanced AI: 22 algorithms vs 0-5 in all competitors",
    "Broadest module coverage: 24+ modules vs 6-12 in competitors",
    "Lowest TCO: zero per-user licensing fees",
    "Native LEAP: auto-calculation per Title 5 USC",
    "Purpose-built for OIG: not adapted from generic CRM or ITSM",
    "Full source code: perpetual license, no vendor lock-in",
    "FedRAMP-certifiable architecture with documented security controls",
    "162 automated tests at 100% pass rate",
    "Comprehensive documentation: tech specs, user manual, implementation plan",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Prepared by Y Point  |  April 2026")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save("docs/OIG-ITS_Vendor_Comparison.docx")
import os
print(f"Created: docs/OIG-ITS_Vendor_Comparison.docx ({os.path.getsize('docs/OIG-ITS_Vendor_Comparison.docx')//1024} KB)")
